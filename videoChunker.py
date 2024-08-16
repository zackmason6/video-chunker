import os
import sys
import pandas as pd
import docx
import csv
#from moviepy.video.io.ffmpeg_tools import ffmpeg_extract_subclip
#from moviepy.editor import VideoFileClip
import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from tkinter import END


def video_operation():
    videoFilePath = videoFileNameEntry.get()
    file_exists = os.path.isfile(videoFilePath)
    if not file_exists:
        messagebox.showinfo("File not found", "The file specified does not exist. Re-enter a file path")
    else:
        #messagebox.showinfo("File found", "This test is successful")
        segment_length = segmentLengthEntry.get()
        #split_video(filename, segment_length)
        messagebox.showinfo("File Found", "Segment length read as: " + str(segment_length))


# Function to handle the submit button click
def submit_data():
    field1_data = entry1.get()
    field2_data = entry2.get()
    field3_data = entry3.get()
    field4_data = entry4.get()
    field5_data = entry5.get()
    field6_data = entry6.get()
    field7_data = entry7.get()
    field8_data = entry8.get()
    field9_data = entry9.get()
    field10_data = entry10.get()
    
    # File name for the CSV
    filename = 'data.csv'
    
    # Check if the file exists to determine whether to write headers
    file_exists = os.path.isfile(filename)
    
    # Open the CSV file for appending data
    with open(filename, mode='a', newline='') as file:
        writer = csv.writer(file)
        
        # Write headers if the file is being created for the first time
        if not file_exists:
            writer.writerow(["Field1", "Field2", "Field3","Field4","Field5","Field6","Field7","Field8","Field9","Field10"])
        
        # Write the data from the text fields
        writer.writerow([field1_data, field2_data, field3_data, field4_data, field5_data, field6_data, field7_data, field8_data, field9_data, field10_data])
    
    # Inform the user that the data has been saved
    messagebox.showinfo("Data Saved", "The data has been successfully saved to the 'data' spreadsheet in your working directory.")
    
    # Clear the entry fields after submission (optional)
    entry1.delete(0, tk.END)
    entry2.delete(0, tk.END)
    entry3.delete(0, tk.END)
    entry4.delete(0, tk.END)
    entry5.delete(0, tk.END)
    entry6.delete(0, tk.END)
    entry7.delete(0, tk.END)
    entry8.delete(0, tk.END)
    entry9.delete(0, tk.END)
    entry10.delete(0, tk.END)

def getFileList(myDirectory,extension):
    """This function gets a list of files from the myDirectory variable.
    Specifically, the os.listdir command gets a list of everything in the directory
    and then this function iterates over that list and grabs everything where the
    filename ends with csv. These are then added to another list.

    Args:
        myDirectory (string): This directory is specified by the user.

    Returns:
        list: This is a list of xml files in the user's current directory.
    """
    dataFiles = []
    myFiles = os.listdir(myDirectory)
    for file in myFiles:
        if file.endswith(extension):
            dataFiles.append(file)
    return dataFiles

#os.environ["IMAGEIO_FFMPEG_EXE"] = r"c:\users\zachary.t.mason\desktop\coding\codingprojects\ncrmp_dataanalysis\.venv\lib\site-packages\ffmpeg"
def startMenu():
    whitelist = [1,2,3]
    print("Please enter the corresponding number of your choice from the menu below:")
    print("1. Split video file into chunks by time")
    print("2. Split video file into chunks by size")
    print("3. Generate metadata")

    userSelection = input("Your selection: ")
    if int(userSelection) not in whitelist:
        print("\nInvalid selection\n")
        startMenu()
    return(userSelection)

def video_upload():
    filename = videoFileNameEntry.get()
    #clip = VideoFileClip(filename)
    #duration = clip.duration
    fileSize = getFileSize(filename)

    video_information_box.delete('1.0',END)
    video_information_box.insert('1.0', "Filename read as: " + str(filename))
    video_information_box.insert('2.0', "\nLength of video: ")
    video_information_box.insert('3.0', "\nSize of video: " + str(fileSize) + "GB")
    #video_information_box.insert('2.0', "Video duration read as: " + str(duration))

def split_video(filename, segment_length):
    segment_length = int(segment_length)
    clip = VideoFileClip(filename)
    duration = clip.duration
    print("Segment length read as: " + str(segment_length))
    print("Filename read as: " + str(filename))
    print("Duration read as: " + str(duration))

    start_time = 0
    end_time = segment_length
    i = 1
    # Extract the filename without extension
    basename = os.path.basename(filename).split('.')[0]

    while start_time < duration:
        # Create output filename based on original filename. Add parameters here?!
        output = os.path.join(f"{basename}_part{i}.mp4")
        print("Output listed as: " + str(output))
        ffmpeg_extract_subclip(filename, start_time, min(end_time, duration), targetname=output)
        start_time = end_time
        end_time += segment_length
        files.download(output)
        i += 1

    print(f'Video split into {i-1} parts.')

def getMetadataFile(listOfFiles):
    metadataFile = input("Which file contains your metadata? ")
    if metadataFile not in listOfFiles:
        print("Invalid entry.\n")
        getMetadataFile(listOfFiles)
    return metadataFile

def getFileSize(file_name):
    file_stats = os.stat(file_name)
    fileSize = file_stats.st_size / (1024 * 1024 * 1024)
    return fileSize

def getData(dataSheet,columnHeader):
  """
  This function performs all of the necessary data gathering from the tracking spreadsheet.
  Additionally, it passes this information to the sendMimeEmail function.
  """
  cellValue = dataSheet.loc[1, columnHeader]
  return cellValue

def readMetadataIn():
    myDirectory = os.getcwd()
    listOfLocalFiles = getFileList(myDirectory,".xlsx")
    print("Here are your local files: ")
    for item in listOfLocalFiles:
        print(str(item))

    metadataFile = getMetadataFile(listOfLocalFiles)
    xlsx = pd.ExcelFile(metadataFile)
    df1 = pd.read_excel(xlsx, sheet_name='Dive Specific Metadata')
    #print(df1)

    myDict = {}

    myColumns = df1.columns
    for column in myColumns:
        tempValue = getData(df1, column)
        tempDict = {column:tempValue}
        myDict.update(tempDict)
        #print(str(column) + ":" + str(tempValue))

    print("Here is the full dictionary: ")
    for key in myDict.keys():
        print(str(key) + ": " + str(myDict.get(key)))

    wordDocList = getFileList(myDirectory,".docx")
    print("Here are your local files: ")
    for item in wordDocList:
        print(str(item))
    wordDocFile = getMetadataFile(wordDocList)
    word = docx.Document(wordDocFile)
    print(str(len(word.paragraphs)))
    for paragraph in word.paragraphs:
        print(str(paragraph.text))

def create_page_content(page):
    # Top frame
    top_frame = tk.Frame(page)
    top_frame.pack(fill=tk.X, padx=10, pady=10)

    # Top label
    top_label = tk.Label(top_frame, text="OER Video Chunking Application", font=('Arial', 16, 'bold'))
    top_label.pack(fill=tk.X)

    # Top description 3
    top_description3 = tk.Label(top_frame, text="This application will split large video files into chunks based on user input and generate a metadata file. Here are the basic directions.",
                                 justify=tk.CENTER, font=('Arial', 10, 'bold'))
    top_description3.bind('<Configure>', lambda e: top_description3.config(wraplength=top_description3.winfo_width()))
    top_description3.pack(pady=5, fill=tk.X, expand=True)

    # Top description 4
    top_description4 = tk.Label(top_frame, text="1. Enter the path to your video file and click submit video to get video information and processing suggestions. \n2. Decide how to split the video based on entered parameters. \n3. Enter any corresponding metadata for your video and generate a metadata spreadsheet.",
                                 justify=tk.LEFT, font=('Arial', 10))
    top_description4.bind('<Configure>', lambda e: top_description4.config(wraplength=top_description4.winfo_width()))
    top_description4.pack(pady=5, fill=tk.X, expand=True)

    # Separator (Bar)
    separator = ttk.Separator(page, orient='horizontal')
    separator.pack(fill=tk.X, padx=10, pady=10)

if __name__ == "__main__":
    userChoice = startMenu()
    print("You selected: " + str(userChoice))

    # Initialize the main application window
    app = tk.Tk()
    app.title("Video Chunker")

    # Center the window on the screen
    app.update_idletasks()  # Ensure correct dimensions are calculated
    window_width = app.winfo_reqwidth()
    window_height = app.winfo_reqheight()
    screen_width = app.winfo_screenwidth()
    screen_height = app.winfo_screenheight()

    x_position = (screen_width // 2) - (window_width // 2)
    y_position = (screen_height // 2) - (window_height // 2)
    app.geometry(f'+{x_position}+{y_position}')

    # Make sure the window pops up in front of other windows
    app.attributes('-topmost', True)
    app.after(1, lambda: app.attributes('-topmost', False))

    # Configure the main window to expand the widgets
    app.columnconfigure(0, weight=1)

    # Create a notebook (for tabbed pages)
    notebook = ttk.Notebook(app)
    notebook.pack(fill=tk.BOTH, expand=True)

    # First page (tab)
    page1 = ttk.Frame(notebook)
    notebook.add(page1, text="Upload Video")

    # Second page (tab)
    page2 = ttk.Frame(notebook)
    notebook.add(page2, text="Split Video")

    # Second page (tab)
    page3 = ttk.Frame(notebook)
    notebook.add(page3, text="Generate Metadata")

    pageList = [page1,page2,page3]

    app.geometry("600x700")

    for page in pageList:
        create_page_content(page)

    ############################# PAGE 1 CONTENT STARTS HERE #################################
    bottom_frame = tk.Frame(page1)
    bottom_frame.pack(fill=tk.X, padx=10, pady=10)

    # Create a frame for the label of the required dimension at the required place
    label_frame = tk.Frame(bottom_frame, width = 200, height = 20)
    label_frame.pack(fill=tk.X,padx=5,pady=5)
    #label_frame.place(x = 60, y = 100)

    # Create the label and pack it to fill the whole frame (default anchor is Centre) 
    test_label = tk.Label(label_frame, text = 'Enter Video Information', font=('Arial', 12, 'bold'))
    test_label.pack(expand = True)

    videoFileName = tk.Label(label_frame, text="Video File Path:", font=('Arial', 10, 'bold'), justify=tk.LEFT, anchor="w")
    videoFileName.pack(pady=5, fill=tk.X, expand=True)
    videoFileNameEntry = tk.Entry(label_frame)
    videoFileNameEntry.bind('<Configure>', lambda e: videoFileNameEntry.config(wraplength=videoFileNameEntry.winfo_width()))
    videoFileNameEntry.pack(pady=5, fill=tk.X, expand=True)
    videoFileNameInstructions = tk.Label(label_frame, text="If this video is in your working directory, just enter the filename. Be sure to include the extension (.mov or .mp4). Get the full file path by finding the file in file explorer, right clicking it and selecting 'properties.'", font=('Arial', 8), justify=tk.LEFT, anchor="w")
    videoFileNameInstructions.bind('<Configure>', lambda e: videoFileNameInstructions.config(wraplength=videoFileNameInstructions.winfo_width()))
    videoFileNameInstructions.pack(pady=5, fill=tk.X, expand=True)

    video_information_box = tk.Text(bottom_frame,height=10)
    video_information_box.bind('<Configure>', lambda e: video_information_box.config(wraplength=video_information_box.winfo_width()))
    video_information_box.pack(pady=5,fill=tk.X, expand =True)
    video_information_box.delete('1.0',END)
    video_information_box.insert('1.0', "Video information and processing suggestions will be displayed here after clicking the upload button below.")

    # Create a submit button
    submit_button = tk.Button(bottom_frame, text="Upload Video", command=video_upload, bg="lightGrey", fg="black")
    submit_button.pack(pady=20)

    ################################### PAGE 2 CONTENT STARTS HERE ############################################

    page2_bottom_frame = tk.Frame(page2)
    page2_bottom_frame.pack(fill=tk.X, padx=10, pady=10)

    # Create a frame for the label of the required dimension at the required place
    page2_label_frame = tk.Frame(page2_bottom_frame, width = 200, height = 20)
    page2_label_frame.pack(fill=tk.X,padx=5,pady=5)
    #label_frame.place(x = 60, y = 100)

    # Create the label and pack it to fill the whole frame (default anchor is Centre) 
    page2_test_label = tk.Label(page2_label_frame, text = 'Enter Chunking Parameters', font=('Arial', 12, 'bold'))
    page2_test_label.pack(expand = True)

    updatedFileName = tk.Label(page2_label_frame, text="Updated Video Name (if applicable): ", font=('Arial', 10, 'bold'), justify=tk.LEFT, anchor="w")
    updatedFileName.pack(pady=5, fill=tk.X, expand=True)
    updatedFileNameEntry = tk.Entry(page2_label_frame)
    updatedFileNameEntry.bind('<Configure>', lambda e: updatedFileNameEntry.config(wraplength=updatedFileNameEntry.winfo_width()))
    updatedFileNameEntry.pack(pady=5, fill=tk.X, expand=True)
    updatedFileNameInstructions = tk.Label(page2_label_frame, text="If you would like to change the name of the output files to something that doesn't match the input filename, enter that here. Otherwise, leave this blank.", font=('Arial', 8), justify=tk.LEFT, anchor="w")
    updatedFileNameInstructions.bind('<Configure>', lambda e: updatedFileNameInstructions.config(wraplength=updatedFileNameInstructions.winfo_width()))
    updatedFileNameInstructions.pack(pady=5, fill=tk.X, expand=True)

    chunkLength = tk.Label(page2_label_frame, text="Desired length of video chunks (in seconds): ", font=('Arial', 10, 'bold'), justify=tk.LEFT, anchor="w")
    chunkLength.pack(pady=5, fill=tk.X, expand=True)
    chunkLengthEntry = tk.Entry(page2_label_frame)
    chunkLengthEntry.bind('<Configure>', lambda e: chunkLengthEntry.config(wraplength=chunkLengthEntry.winfo_width()))
    chunkLengthEntry.pack(pady=5, fill=tk.X, expand=True)
    chunkLengthInstructions = tk.Label(page2_label_frame, text="Enter the length (in seconds) that you would like each video chunk to be. Note that the last chunk may be shorter than the rest.", font=('Arial', 8), justify=tk.LEFT, anchor="w")
    chunkLengthInstructions.bind('<Configure>', lambda e: chunkLengthInstructions.config(wraplength=chunkLengthInstructions.winfo_width()))
    chunkLengthInstructions.pack(pady=5, fill=tk.X, expand=True)

    split_video_button = tk.Button(page2, text="Split Video", command=video_operation, bg="lightGrey", fg="black")
    split_video_button.pack(pady=20)

    # Bold separator (using a Frame)
    #separator_frame = tk.Frame(page2, height=3, bg="black")
    #separator_frame.pack(fill=tk.X, padx=10, pady=10)

    ############################### PAGE 3 STARTS HERE #############################

    # Create a frame that will contain the canvas and the scrollbar
    outer_frame = tk.Frame(page3)
    outer_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    # Create a canvas inside the outer_frame
    canvas = tk.Canvas(outer_frame)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Add a scrollbar to the canvas
    scrollbar = tk.Scrollbar(outer_frame, orient="vertical", command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill="y")

    # Configure the canvas to work with the scrollbar
    canvas.configure(yscrollcommand=scrollbar.set)

    # Create a frame inside the canvas
    label_frame = tk.Frame(canvas)

    # Add the label_frame to the canvas
    window = canvas.create_window((0, 0), window=label_frame, anchor="nw")

    # Adjust the size of the label_frame to match the outer_frame's width
    def on_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))
        canvas.itemconfig(window, width=canvas.winfo_width())

    label_frame.bind("<Configure>", on_configure)

    # Populate the label_frame with widgets
    test_label = tk.Label(label_frame, text='Enter Video Metadata', font=('Arial', 10, 'bold'))
    test_label.pack(expand=True)

    # Adjust the canvas size when the outer frame is resized
    outer_frame.bind('<Configure>', lambda event: canvas.itemconfig(window, width=event.width))

    # Enable mouse wheel scrolling
    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    # Bind the mousewheel event to the canvas
    canvas.bind_all("<MouseWheel>", _on_mousewheel)

    label1 = tk.Label(label_frame, text="Field 1:", justify=tk.LEFT, anchor="w")
    label1.pack(pady=5, fill=tk.X, expand=True)
    entry1 = tk.Entry(label_frame)
    entry1.pack(pady=5, fill=tk.X, expand=True)

    label2 = tk.Label(label_frame, text="Field 2:", justify=tk.LEFT, anchor="w")
    label2.pack(pady=5, fill=tk.X, expand=True)
    entry2 = tk.Entry(label_frame)
    entry2.pack(pady=5, fill=tk.X, expand=True)

    label3 = tk.Label(label_frame, text="Field 3:", justify=tk.LEFT, anchor="w")
    label3.pack(pady=5, fill=tk.X, expand=True)
    entry3 = tk.Entry(label_frame)
    entry3.pack(pady=5, fill=tk.X, expand=True)

    label4 = tk.Label(label_frame, text="Field 4:", justify=tk.LEFT, anchor="w")
    label4.pack(pady=5, fill=tk.X, expand=True)
    entry4 = tk.Entry(label_frame)
    entry4.pack(pady=5, fill=tk.X, expand=True)

    label5 = tk.Label(label_frame, text="Field 5:", justify=tk.LEFT, anchor="w")
    label5.pack(pady=5, fill=tk.X, expand=True)
    entry5 = tk.Entry(label_frame)
    entry5.pack(pady=5, fill=tk.X, expand=True)

    label6 = tk.Label(label_frame, text="Field 6:", justify=tk.LEFT, anchor="w")
    label6.pack(pady=5, fill=tk.X, expand=True)
    entry6 = tk.Entry(label_frame)
    entry6.pack(pady=5, fill=tk.X, expand=True)

    label7 = tk.Label(label_frame, text="Field 7:", justify=tk.LEFT, anchor="w")
    label7.pack(pady=5, fill=tk.X, expand=True)
    entry7 = tk.Entry(label_frame)
    entry7.pack(pady=5, fill=tk.X, expand=True)

    label8 = tk.Label(label_frame, text="Field 8:", justify=tk.LEFT, anchor="w")
    label8.pack(pady=5, fill=tk.X, expand=True)
    entry8 = tk.Entry(label_frame)
    entry8.pack(pady=5, fill=tk.X, expand=True)

    label9 = tk.Label(label_frame, text="Field 9:", justify=tk.LEFT, anchor="w")
    label9.pack(pady=5, fill=tk.X, expand=True)
    entry9 = tk.Entry(label_frame)
    entry9.pack(pady=5, fill=tk.X, expand=True)

    label10 = tk.Label(label_frame, text="Field 10:", justify=tk.LEFT, anchor="w")
    label10.pack(pady=5, fill=tk.X, expand=True)
    entry10 = tk.Entry(label_frame)
    entry10.pack(pady=5, fill=tk.X, expand=True)
    # Add more widgets to the label_frame if needed

    # Update the scrollregion after adding all the widgets
    label_frame.update_idletasks()
    canvas.config(scrollregion=canvas.bbox("all"))

    # Create a submit button
    submit_button = tk.Button(page3, text="Generate Metadata", command=submit_data, bg="lightGrey", fg="black")
    submit_button.pack(pady=20)

    # Run the application
    app.mainloop()

    #cruiseID = getData(df1,'Project ID (or CruiseID)')
    #diveID = getData(df1,'DiveID')
    #print("HERE IS WHAT YOU EXTRACTED FROM THE FILE: " + str(testString))

    #cruise = input("Enter the cruise name: ")
    #diveSite = input("Enter the dive site name: ")

    #video_path = r'C:\Users\Zachary.T.Mason\Desktop\PS2220_DIVE0982_20220504_MOHAWKROVHD.mov'
    #segment_length = int(segment_length)
    #segment_length = 60
    #output_dir = sys.argv[3]  # third argument from command line
    #split_video(video_path, segment_length)